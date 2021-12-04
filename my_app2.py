from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

document.add_picture('moi.png', width=Inches(2.5))

speak('was ist deine Name? ')
name = input('What is your name? ')

speak('wie lautet dein telephon nummer? ')
phone_number = input('What is your phone muber? ')

email =input('What is your email adresse? ') 

document.add_paragraph(name +'\n' + phone_number +'\n'+ email+'\n')

# about me
document.add_heading('Work experoence')
document.add_paragraph(input('Tell me about yourself? '))

# work experience
document.add_heading('Work experience')
p = document.add_paragraph()

company = input('enter company name: ')
from_date = input('from Date: ')
to_date = input('to Date: ')

p.add_run(company + ': ').bold = True
p.add_run(from_date + '->' + to_date +'\n').italic =True

experience_details = input('Describe your experince at: ' +company+'\n')
p.add_run(experience_details)

# more experiences
document.add_heading('Mores experiences')
while True:
    has_more_experinces = input('Do u have more experinces? Yes or No:  ' )
    if has_more_experinces.lower() == 'yes':
        p = document.add_paragraph()

        company = input('enter company name: ')
        from_date = input('from Date: ')
        to_date = input('to Date: ')

        p.add_run(company + ': ').bold = True
        p.add_run(from_date + '-' + to_date +'\n').italic =True

        experience_details = input('Describe your experince at: ' +company+'\n')
        p.add_run(experience_details)

    else:
        break


# Skills

document.add_heading('Skills')
#skill = input('Enter skill')
#p = document.add_paragraph(skill)
#p.style = 'List Bullet'
while True:
    has_skills = input('Do you have skills? Yes or No: ')
    if has_skills.lower() == 'yes' or  has_skills.lower() == 'y':
       
        speak('your skills: ')
        skills = input('your skills: ')
        
        p = document.add_paragraph(skills)
        p.style = 'List Bullet'
        
    else:
        break

# footer

section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using kokoeaes: " +from_date
document.save('cv.docx')
